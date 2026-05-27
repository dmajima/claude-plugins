"""構造化議事録データ（JSON v2.0）から docx ファイルを生成する"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
sys.stderr.reconfigure(encoding='utf-8')

import argparse
import json
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

COLOR_PRIMARY = RGBColor(0x1B, 0x3A, 0x5C)
COLOR_HEADER_BG = 'E8EEF4'
COLOR_ACCENT_BG = 'F0F4F8'
COLOR_WHITE = 'FFFFFF'
FONT_NAME = 'Meiryo'


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = FONT_NAME
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), FONT_NAME)
        rFonts.set(qn('w:hAnsi'), FONT_NAME)
        rFonts.set(qn('w:eastAsia'), FONT_NAME)
    return p


def set_table_col_widths(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    layout = tbl_pr.find(qn('w:tblLayout'))
    if layout is None:
        layout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
        tbl_pr.append(layout)
    else:
        layout.set(qn('w:type'), 'fixed')

    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                tc = row.cells[idx]._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = parse_xml(f'<w:tcW {nsdecls("w")}/>')
                    tcPr.insert(0, tcW)
                tcW.set(qn('w:w'), str(int(width.emu / 635)))
                tcW.set(qn('w:type'), 'dxa')


def set_cell_shading(cell, color_hex):
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell(cell, text, bold=False, font_size=Pt(10), color=None, align=None):
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold
    run.font.size = font_size
    run.font.name = FONT_NAME
    if color:
        run.font.color.rgb = color
    if align:
        cell.paragraphs[0].alignment = align


def add_title_section(doc, metadata):
    title_text = metadata.get('title', '会議議事録')
    p_title = doc.add_paragraph()
    p_title.style = doc.styles['Title']
    run = p_title.add_run(f'【会議議事録】')
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_PRIMARY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_sub = p_sub.add_run(title_text)
    run_sub.font.size = Pt(18)
    run_sub.font.bold = True
    run_sub.font.name = FONT_NAME
    run_sub.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph()


def add_meta_table(doc, metadata):
    date_str = metadata.get('date', '')
    start = metadata.get('startTime', '')
    end = metadata.get('endTime', '')
    dur = metadata.get('durationMinutes', '')
    time_display = f"{date_str}  {start} - {end}" if start else date_str
    if dur:
        time_display += f"（{dur}分）"

    location = metadata.get('location', 'オンライン会議')
    created_by = metadata.get('createdBy', 'AI（文字起こし + Claude 構造化）')

    participants = metadata.get('participants', [])
    host_orgs = {p.get('organization', '') for p in participants if p.get('role') == 'host'}
    by_org = {}
    for p in participants:
        org = p.get('organization', '')
        name = p.get('name', '')
        if org not in by_org:
            by_org[org] = []
        by_org[org].append(name)

    att_lines = []
    for org, names in by_org.items():
        prefix = f"{org} / " if org else ""
        for n in names:
            suffix = "様" if org and org not in host_orgs else ""
            att_lines.append(f"{prefix}{n}{suffix}")
    att_text = '\n'.join(att_lines) if att_lines else ''

    rows_data = [
        ('日時', time_display),
        ('場所', location),
        ('出席者', att_text),
        ('議事録作成', created_by),
    ]

    table = doc.add_table(rows=len(rows_data), cols=2, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_col_widths(table, [Cm(2.2), Cm(13.8)])

    for i, (label, value) in enumerate(rows_data):
        cell_label = table.cell(i, 0)
        cell_value = table.cell(i, 1)

        set_cell_shading(cell_label, COLOR_HEADER_BG)
        set_cell(cell_label, label, bold=True, font_size=Pt(9.5), color=COLOR_PRIMARY)

        set_cell(cell_value, value, font_size=Pt(9.5))

    doc.add_paragraph()


def add_agenda_toc(doc, agendas):
    add_heading(doc, '議題一覧', level=1)
    for a in agendas:
        doc.add_paragraph(a['title'], style='List Number')
    doc.add_paragraph()


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = FONT_NAME
    run.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)


def add_bullet_list(doc, items):
    for item in items:
        doc.add_paragraph(str(item), style='List Bullet')


def add_agenda_section(doc, agenda):
    add_heading(doc, f"{agenda['id']}. {agenda['title']}", level=1)

    bg = agenda.get('background', '')
    if bg:
        add_section_heading(doc, '背景・目的:')
        doc.add_paragraph(bg)

    specs = agenda.get('specifications', [])
    if specs:
        add_section_heading(doc, '仕様・機能詳細:')
        add_bullet_list(doc, specs)

    discussions = agenda.get('discussions', [])
    if discussions:
        add_section_heading(doc, '議論の内容:')
        add_bullet_list(doc, discussions)

    concerns = agenda.get('concerns', [])
    if concerns:
        add_section_heading(doc, '懸念点:')
        add_bullet_list(doc, concerns)

    conclusions = agenda.get('conclusions', [])
    if conclusions:
        add_section_heading(doc, '結論・合意事項:')
        add_bullet_list(doc, conclusions)


def add_action_items(doc, items):
    add_heading(doc, 'アクションまとめ', level=1)
    if not items:
        doc.add_paragraph('なし')
        return

    table = doc.add_table(rows=1 + len(items), cols=4, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_col_widths(table, [Cm(0.8), Cm(2.5), Cm(1.8), Cm(10.9)])

    headers = ['#', '担当', '期限', '内容']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, COLOR_HEADER_BG)
        set_cell(cell, h, bold=True, font_size=Pt(9), color=COLOR_PRIMARY,
                 align=WD_ALIGN_PARAGRAPH.CENTER)

    for row_idx, item in enumerate(items, 1):
        set_cell(table.cell(row_idx, 0), item.get('id', row_idx), font_size=Pt(9),
                 align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(table.cell(row_idx, 1), item.get('assignee', ''), font_size=Pt(9))
        set_cell(table.cell(row_idx, 2), item.get('deadline', ''), font_size=Pt(9))
        set_cell(table.cell(row_idx, 3), item.get('content', ''), font_size=Pt(9))

        if row_idx % 2 == 0:
            for col in range(4):
                set_cell_shading(table.cell(row_idx, col), COLOR_ACCENT_BG)

    doc.add_paragraph()


def add_next_meeting(doc, next_meeting):
    add_heading(doc, '次回予定', level=1)
    if not next_meeting:
        doc.add_paragraph('未定')
        return

    rows_data = []
    if next_meeting.get('date'):
        rows_data.append(('日時', next_meeting['date']))
    agendas = next_meeting.get('plannedAgendas', [])
    if agendas:
        rows_data.append(('議題（予定）', ', '.join(agendas)))
    preps = next_meeting.get('preparations', [])
    if preps:
        prep_text = '\n'.join(
            f"- {p.get('task', '')}（{p.get('assignee', '')}）" for p in preps
        )
        rows_data.append(('準備事項', prep_text))

    if rows_data:
        table = doc.add_table(rows=len(rows_data), cols=2, style='Table Grid')
        set_table_col_widths(table, [Cm(2.8), Cm(13.2)])
        for i, (label, value) in enumerate(rows_data):
            set_cell_shading(table.cell(i, 0), COLOR_HEADER_BG)
            set_cell(table.cell(i, 0), label, bold=True, font_size=Pt(9), color=COLOR_PRIMARY)
            set_cell(table.cell(i, 1), value, font_size=Pt(9))


def generate(input_path: str, template_path: Optional[str], output_path: str):
    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except FileNotFoundError:
        print(f"ERROR: Input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"ERROR: Invalid JSON in {input_path}: {e}", file=sys.stderr)
        sys.exit(1)

    if template_path and Path(template_path).exists():
        doc = Document(template_path)
    else:
        doc = Document()

    metadata = data.get('metadata', {})
    add_title_section(doc, metadata)
    add_meta_table(doc, metadata)

    agendas = data.get('agendas', [])
    if agendas:
        add_agenda_toc(doc, agendas)
        for agenda in agendas:
            add_agenda_section(doc, agenda)

    add_action_items(doc, data.get('actionItems', []))
    add_next_meeting(doc, data.get('nextMeeting'))

    doc.save(output_path)
    print(f"Generated: {output_path}")


def main():
    parser = argparse.ArgumentParser(description='Generate meeting minutes docx')
    parser.add_argument('--input', required=True, help='Path to minutes.json')
    parser.add_argument('--template', default=None, help='Path to template.docx')
    parser.add_argument('--output', required=True, help='Output docx path')
    args = parser.parse_args()
    generate(args.input, args.template, args.output)


if __name__ == '__main__':
    main()
